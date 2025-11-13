from docx import Document

doc = Document()

text = """Technical Skills
• Languages: Python, Java, C#
• Scripting Languages: Bash, Python scripting, VBscripting
• Query Languages: SQL, HiveQL, Spark SQL
• Frameworks/Tools: Apache Spark, Hadoop (HDFS, MapReduce, Hive), Apache Airflow, Apache Beam, Kafka, TensorFlow, Git, Informatica
• Cloud Technologies: Microsoft Azure (Data Lake Storage, Databricks, Azure SQL), AWS (S3, EMR, Redshift, Glue, Athena)
• Databases: PostgreSQL, MySQL, SQL Server, MongoDB
• IDEs: PyCharm, Visual Studio Code, JupyterLab
• Operating Systems: Linux (Ubuntu), Windows
• Development Tools: Git, GitHub
• Data Visualization: Tableau, Matplotlib, Seaborn, Plotly Express

Certifications
• Certified AWS AI Practitioner

Employment Experience
Schlumberger | Data Engineer | India | 2019–2023
• Designed and implemented scalable batch and streaming data pipelines on Azure and AWS, processing seismic and drilling data, improving data availability to 99.5%.
• Migrated legacy on-prem ETL systems to cloud-native architecture using Apache Spark, Glue, and EMR, reducing operational costs by 20% and accelerating data delivery.
• Automated complex data validation frameworks using Apache Beam and Airflow, reducing manual QA efforts by 80% and ensuring high-quality, reliable data for downstream analytics.
• Utilized ETL tools including Informatica/Pentaho and Airflow for orchestrating complex data workflows, enhancing data pipeline reliability and maintainability.
• Developed and optimized data models and SQL data stores (PostgreSQL), improving query performance by up to 40%.
• Integrated OSDU data standards to enhance cross-domain analytics and improve data ingestion efficiency for global subsurface teams.
• Collaborated with data scientists and software engineers to prepare clean, production-ready data sets supporting ML models and advanced analytics.
• Built and maintained executive-level dashboards using Tableau, enabling real-time insights into pipeline health and operational KPIs.
• Developed executive dashboards in Tableau, providing real-time insights into pipeline health, ingestion volume, and SLA performance.
Tech Stack: Hadoop, Spark, Hive, Pig, HDFS, Python, SQL, Tableau, Dataflow, Azure, Azure Data Lake, Apache Beam, OSDU, PostgreSQL, Tableau, Informatica.

Academic and Technical Research Projects
Generative RAG: Retrieval Summarization System on CNN Dataset | University of North Texas | 2024
• Engineered an AI-driven Retrieval-Augmented Generation (RAG) system using Python, NumPy, Pandas, Scikit-learn, and Hugging Face Transformers to deliver advanced question answering and text summarization capabilities.
• Integrated SpaCy and Gensim for natural language processing, achieving a 35% improvement in response accuracy over baseline extractive methods.
• Designed and deployed a user-friendly front-end using Streamlit. Implemented robust evaluation metrics and error analysis workflows.

EV Cars Statistical Dashboard – React Application | University of North Texas | 2024
• Developed a dynamic interactive dashboard using React, Node.js, and Nivo to visualize electric vehicle statistics from large-scale GeoJSON datasets.
• Leveraged Material-UI (MUI) and D3.js to create highly engaging and responsive visualizations.
• Optimized rendering and state management strategies, reducing dashboard load times by 20% and improving overall user experience.

Big Data and Data Science Project: Customer segmentation ML | University of North Texas | 2023
• Designed and deployed segmentation models using PySpark and MLlib to drive marketing analytics and customer insights.
• Integrated Cloudera Hadoop ecosystem and PySpark ecosystem. Performed data processing and querying in Cloudera and machine learning in Pyspark locally.
"""

for line in text.split("\n"):
    if line.strip().startswith("•"):
        p = doc.add_paragraph()
        p.add_run(line.strip()).bold = False
    else:
        doc.add_heading(line, level=2)


doc.save("doc.docx")


